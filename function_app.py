import azure.functions as func
import logging
import requests
from pypdf import PdfReader
from docx import Document
from io import BytesIO

app = func.FunctionApp(http_auth_level=func.AuthLevel.ADMIN)

@app.route(route="doc_to_text")
def doc_to_text(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

    # 1. Parse the Request Body
    try:
        req_body = req.get_json()
        doc_url = req_body.get('url')
    except ValueError:
        return func.HttpResponse(
            "Invalid JSON. Please pass a body like {'url': 'http://example.com/file.pdf'} or {'url': 'http://example.com/file.docx'}",
            status_code=400
        )

    if not doc_url:
        return func.HttpResponse(
            "Please pass a 'url' in the request body.",
            status_code=400
        )

    # 2. Download the document
    try:
        response = requests.get(doc_url, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        return func.HttpResponse(f"Error downloading URL: {str(e)}", status_code=400)

    # 3. Verify Content Type (Optional but recommended)
    content_type = response.headers.get('Content-Type', '')
    is_pdf = 'application/pdf' in content_type or doc_url.lower().endswith('.pdf')
    is_docx = ('application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type or 
               doc_url.lower().endswith('.docx'))
    
    if not (is_pdf or is_docx):
         return func.HttpResponse("URL does not appear to be a supported document file (PDF or DOCX).", status_code=400)

    # 4. Extract Text
    try:
        # Wrap the binary content in a BytesIO stream
        f = BytesIO(response.content)
        
        if is_pdf:
            # Extract text from PDF
            reader = PdfReader(f)
            extracted_text = []
            
            # Iterate over pages
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    extracted_text.append(text)
                else:
                    extracted_text.append(f"[Could not extract text from page {page_num + 1}]")
            
            full_text = "\n\n".join(extracted_text)
            
        elif is_docx:
            # Extract text from Word document
            doc = Document(f)
            extracted_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text.append(" | ".join(row_text))
            
            full_text = "\n\n".join(extracted_text)
        
        else:
            return func.HttpResponse("Unsupported document type.", status_code=400)

        # Return the text as a JSON response
        return func.HttpResponse(
            full_text,
            mimetype="text/plain",
            status_code=200
        )

    except Exception as e:
        logging.error(f"Error parsing document: {e}")
        return func.HttpResponse(f"Error parsing document data: {str(e)}", status_code=500)